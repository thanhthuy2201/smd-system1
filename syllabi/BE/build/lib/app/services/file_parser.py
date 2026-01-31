"""File parsing service for syllabus documents (PDF, DOCX, TXT)"""
import re
import logging
from typing import Optional
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ParsedSyllabus:
    """Parsed syllabus content"""
    title: Optional[str] = None
    course_code: Optional[str] = None
    course_name: Optional[str] = None
    credits: Optional[int] = None
    total_hours: Optional[int] = None
    academic_year: Optional[str] = None
    semester: Optional[str] = None
    prerequisites: Optional[str] = None
    learning_outcomes: Optional[str] = None
    assessment_methods: Optional[str] = None
    textbooks: Optional[str] = None
    teaching_methods: Optional[str] = None
    materials: Optional[str] = None
    raw_text: str = ""


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        import pypdf
        from io import BytesIO

        reader = pypdf.PdfReader(BytesIO(file_content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except ImportError:
        logger.warning("pypdf not installed. Install with: pip install pypdf")
        raise ValueError("PDF processing not available. pypdf package required.")
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
        raise ValueError(f"Failed to extract PDF content: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"

        return text
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        raise ValueError("DOCX processing not available. python-docx package required.")
    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {e}")
        raise ValueError(f"Failed to extract DOCX content: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Unable to decode text file")
    except Exception as e:
        logger.error(f"Failed to read TXT file: {e}")
        raise ValueError(f"Failed to read text file: {str(e)}")


def extract_text(file_content: bytes, filename: str) -> str:
    """Extract text from file based on extension"""
    ext = Path(filename).suffix.lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_content)
    elif ext == '.docx':
        return extract_text_from_docx(file_content)
    elif ext == '.txt':
        return extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")


def parse_syllabus_content(text: str) -> ParsedSyllabus:
    """
    Parse extracted text to identify syllabus components.
    Uses regex patterns to find common syllabus sections.
    """
    result = ParsedSyllabus(raw_text=text)

    # Normalize text
    text_lower = text.lower()
    lines = text.split('\n')

    # Extract course code (patterns like CS101, MATH-201, etc.)
    code_pattern = r'\b([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)\b'
    code_match = re.search(code_pattern, text)
    if code_match:
        result.course_code = code_match.group(1).replace(' ', '')

    # Extract credits (patterns like "3 credits", "Credits: 3", etc.)
    credits_patterns = [
        r'credits?\s*[:=]?\s*(\d+)',
        r'(\d+)\s*credits?',
        r'số tín chỉ\s*[:=]?\s*(\d+)',  # Vietnamese
        r'tín chỉ\s*[:=]?\s*(\d+)',
    ]
    for pattern in credits_patterns:
        match = re.search(pattern, text_lower)
        if match:
            result.credits = int(match.group(1))
            break

    # Extract total hours
    hours_patterns = [
        r'total\s*hours?\s*[:=]?\s*(\d+)',
        r'(\d+)\s*hours?\s*total',
        r'contact\s*hours?\s*[:=]?\s*(\d+)',
        r'số tiết\s*[:=]?\s*(\d+)',  # Vietnamese
    ]
    for pattern in hours_patterns:
        match = re.search(pattern, text_lower)
        if match:
            result.total_hours = int(match.group(1))
            break

    # Extract academic year (pattern like 2024-2025)
    year_pattern = r'(20\d{2}[-/]20\d{2})'
    year_match = re.search(year_pattern, text)
    if year_match:
        result.academic_year = year_match.group(1).replace('/', '-')

    # Extract semester
    semester_patterns = [
        r'semester\s*[:=]?\s*(fall|spring|summer|1|2|i|ii)',
        r'(fall|spring|summer)\s*semester',
        r'học kỳ\s*[:=]?\s*(\d|i|ii)',  # Vietnamese
    ]
    for pattern in semester_patterns:
        match = re.search(pattern, text_lower)
        if match:
            sem = match.group(1).capitalize()
            # Normalize semester names
            sem_map = {'1': 'Fall', '2': 'Spring', 'I': 'Fall', 'Ii': 'Spring'}
            result.semester = sem_map.get(sem, sem)
            break

    # Extract sections based on common headers
    sections = {
        'prerequisites': [
            r'prerequisite[s]?\s*[:=]?\s*(.+?)(?=\n\n|\n[A-Z]|$)',
            r'điều kiện tiên quyết\s*[:=]?\s*(.+?)(?=\n\n|\n[A-Z]|$)',
        ],
        'learning_outcomes': [
            r'(?:course\s*)?learning\s*outcome[s]?\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|\nassessment|\ngrading|$)',
            r'clo\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
            r'mục tiêu\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
            r'chuẩn đầu ra\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
        ],
        'assessment_methods': [
            r'(?:assessment|grading|evaluation)\s*(?:method[s]?|criteria|breakdown)?\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|\ntextbook|\nreference|$)',
            r'đánh giá\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
        ],
        'textbooks': [
            r'(?:textbook[s]?|reference[s]?|required\s*reading)\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|\nteaching|$)',
            r'tài liệu\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
            r'giáo trình\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
        ],
        'teaching_methods': [
            r'teaching\s*method[s]?\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
            r'instructional\s*(?:method[s]?|approach)\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
            r'phương pháp giảng dạy\s*[:=]?\s*(.+?)(?=\n\n[A-Z]|$)',
        ],
    }

    for field, patterns in sections.items():
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                content = match.group(1).strip()
                # Clean up the content
                content = re.sub(r'\s+', ' ', content)
                if len(content) > 10:  # Minimum content length
                    setattr(result, field, content[:2000])  # Limit length
                    break

    return result


async def process_syllabus_file(file_content: bytes, filename: str) -> ParsedSyllabus:
    """
    Process uploaded syllabus file and extract structured content.

    Args:
        file_content: Raw file bytes
        filename: Original filename with extension

    Returns:
        ParsedSyllabus with extracted content
    """
    # Extract text from file
    text = extract_text(file_content, filename)

    if not text or len(text.strip()) < 50:
        raise ValueError("File appears to be empty or contains too little text")

    # Parse the content
    parsed = parse_syllabus_content(text)

    return parsed


def get_supported_extensions() -> list[str]:
    """Get list of supported file extensions"""
    return ['.pdf', '.docx', '.txt']


def validate_file_extension(filename: str) -> bool:
    """Check if file extension is supported"""
    ext = Path(filename).suffix.lower()
    return ext in get_supported_extensions()
